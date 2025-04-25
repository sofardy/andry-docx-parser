import os
import re
import csv
import datetime
from docx import Document

# ...existing code (if any)...
BASE_DIR = os.path.join(os.path.dirname(__file__), 'files')
OUTPUT_CSV = os.path.join(os.path.dirname(__file__), 'output.csv')
YEAR = datetime.datetime.now().year

email_pattern = re.compile(r"[\w\.-]+@[\w\.-]+\.\w+")
phone_pattern = re.compile(r"\d{3}-\d{3}-\d{4}")

def extract_emails(docx_path):
    try:
        doc = Document(docx_path)
    except Exception:
        return []
    texts = []
    for para in doc.paragraphs:
        texts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    content = '\n'.join(texts)
    return email_pattern.findall(content)

def parse_name_phone(filename):
    name_phone = os.path.splitext(filename)[0]
    name_phone = re.sub(r'^\([^)]*\)\s*', '', name_phone)
    match = phone_pattern.search(name_phone)
    phone = match.group() if match else ''
    name = name_phone[:match.start()].strip() if match else name_phone.strip()
    return name, phone

def main():
    rows = []
    for root, _, files in os.walk(BASE_DIR):
        for file in files:
            if file.lower().endswith('.docx'):
                path = os.path.join(root, file)
                rel_folder = os.path.relpath(root, BASE_DIR)
                emails = extract_emails(path)
                name, phone = parse_name_phone(file)
                if not emails:
                    rows.append([file, name, '', phone, rel_folder, YEAR])
                else:
                    for email in emails:
                        rows.append([file, name, email, phone, rel_folder, YEAR])
    with open(OUTPUT_CSV, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow(['file name', 'name', 'email', 'phone', 'folder', 'year'])
        writer.writerows(rows)

if __name__ == '__main__':
    main()
